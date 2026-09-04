# -*- coding: utf-8 -*-
"""在 docx 的关键词段落后插入分页符，使「摘要 + 关键词」独占一页。

用法：
    python fix_docx_pagebreak.py <目标.docx>

要点：
- 路径必须走命令行参数，**不要硬编码** —— 主 docx 常被 Windows 索引锁住，
  实际输出文件名可能是 `_论文_定稿版_新.docx`。
- 幂等：已存在 type=page 的 w:br 时直接跳过，重复运行不会插第二个分页符。
- 找不到「关键词：」段落时明确报错并以非 0 退出，方便流水线兜底检查。
"""
import sys
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def already_has_pagebreak(doc):
    for p in doc.paragraphs:
        if p._p.findall(".//{%s}br[@%s='page']" % (W, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type")):
            return True
    return False


def main(path):
    doc = Document(path)

    if already_has_pagebreak(doc):
        print("已有分页符，跳过")
        return 0

    for para in doc.paragraphs:
        if "关键词" in para.text:
            run = para.add_run()
            br = OxmlElement("w:br")
            br.set(qn("w:type"), "page")
            run._r.append(br)
            doc.save(path)
            print("已在「%s…」后插入分页符" % para.text[:20])
            return 0

    print("❌ 未找到关键词段落，未插入分页符")
    return 1


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(2)
    sys.exit(main(sys.argv[1]))
