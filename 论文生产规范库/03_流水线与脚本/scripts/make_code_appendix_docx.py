# -*- coding: utf-8 -*-
"""按 code-appendix-word skill 生成代码附录 Word。

每文件 = 一个 1x1 Table Grid 边框表格，多色语法高亮（K/C/S/N/F 五色）；
注释清理、留白、字号/行距等均按 skill 落地。
"""
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(BASE, "附录_全部Python代码.docx")

FILES = [
    ('solver.py',  'CA-IGA 求解器：数据读取、气候感知解码器、协同进化遗传算法'),
    ('run_all.py', '主流程：问题1/2 求解、敏感性分析、问题3 动态重调度、绘图与报告'),
]

# ---- 配色 ----
K = RGBColor(0, 0, 170)
C = RGBColor(0, 128, 0)
S = RGBColor(178, 34, 34)
N = RGBColor(138, 43, 226)
F = RGBColor(0, 128, 128)
BLK = RGBColor(0, 0, 0)

# ---- Python 关键字 / 内置 ----
PY_KW = {'and', 'as', 'assert', 'break', 'class', 'continue', 'def', 'del',
         'elif', 'else', 'except', 'finally', 'for', 'from', 'global', 'if',
         'import', 'in', 'is', 'lambda', 'not', 'or', 'pass', 'raise',
         'return', 'try', 'while', 'with', 'yield',
         'None', 'True', 'False'}
PY_BI = {'print', 'len', 'range', 'int', 'float', 'str', 'list', 'dict',
         'set', 'max', 'min', 'sum', 'abs', 'sorted', 'zip', 'enumerate',
         'open', 'round', 'type', 'isinstance', 'any', 'all', 'tuple',
         'map', 'filter', 'reversed', 'hasattr', 'getattr', 'setattr',
         'next', 'iter', 'super', 'object', 'staticmethod', 'classmethod',
         'property', 'globals', 'locals', 'dir', 'vars', 'hex', 'oct',
         'bin', 'pow'}

# ---- 注释删除关键词 ----
DROP_KEYWORDS = (
    '教学', '调试', '解释', '费曼', '帮你理解', '理解', '数据单', '来源',
    '见solve', '为何这么写', '我们', '大家', '即可', '不难', '请看', '参见',
    '目的', '背景', '注释', '备忘', '约定', '说明', '待确认',
    'BUG', 'FIXME', 'TODO', 'XXX', 'HACK', '兼容', '作者', 'Author',
    'code-hand', 'sub-agent', '并行工作流', 'Day 2', '完整初稿', '定稿示例',
    '正式参赛', '据实补全', '本模拟', '本文件', '练习', '对话', '本节',
    '小羽', 'WorkBuddy', 'AI 工具', 'AI工具', '队友', '教练', '指导',
    '评审', '题意', '题目', '论文草稿', '草稿', '仅需', '思路',
    '原值', '记忆', '记录', '复盘', '验证',
)
DROP_PREFIX = (
    '# ---', '# ===', '# __', '#***', '# 重要', '# 关键思路', '# 备忘',
    '# 自我纠错', '# 流程演练',
)
MAX_COMMENT_LEN = 80
MAX_DOCSTRING_LEN = 220
SPECIAL_PATTERN = re.compile(r'[★☆♥♦♣♠✗✓✘✚✦◆■▲△▼▽◯◎●]')


def _is_separator(s):
    return bool(re.match(r'^#[\s\-=_*#·•]+$', s))


def _has_special(s):
    return bool(SPECIAL_PATTERN.search(s))


def clean(lines):
    """精简注释、删除叙事/角色/演练/特殊符号相关内容。"""
    out, prev_blank = [], False
    for raw in lines:
        s = raw.strip()

        # 1) 单行内含三引号（成对闭合）：太长或含 DROP_KEYWORDS 则整段删除
        if '"""' in s or "'''" in s:
            dd = s.count('"""'); ss = s.count("'''")
            if dd == 2 or ss == 2:
                if (len(s) > MAX_DOCSTRING_LEN
                        or any(k in s for k in DROP_KEYWORDS)):
                    continue
                out.append(raw.rstrip()); prev_blank = False
                continue
            # 多行 docstring 的中间行：保留原文（python 解析需要）
            out.append(raw.rstrip()); prev_blank = False
            continue

        # 2) 空行
        if not s:
            if not prev_blank:
                out.append('')
            prev_blank = True
            continue
        prev_blank = False

        # 3) 注释
        if s.startswith('#'):
            if _is_separator(s):
                continue
            if any(s.startswith(p) for p in DROP_PREFIX):
                continue
            if any(k in s for k in DROP_KEYWORDS):
                continue
            if len(s) > MAX_COMMENT_LEN:
                continue
            if _has_special(s):
                continue
            out.append(raw.rstrip())
            continue

        # 4) 代码
        out.append(raw.rstrip())
    return out


def add_run(p, text, color, size=8.5, bold=False, italic=False, font='Consolas'):
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.bold = bold
    r.italic = italic
    r.font.name = font
    rPr = r._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font)
    rFonts.set(qn('w:ascii'), font)
    rFonts.set(qn('w:hAnsi'), font)


# 顺序很关键：注释 > docstring > 字符串 > 装饰器 > 数字 > 标识符
TOKENS = [
    (re.compile(r'#[^\n]*'),                                  C, False, True),
    (re.compile(r'(\"\"\"[\s\S]*?\"\"\"|\'\'\'[\s\S]*?\'\'\')'),
                                                            C, False, True),
    (re.compile(r'(\"[^\"\n]*\"|\'[^\'\n]*\')'),             S, False, False),
    (re.compile(r'(@\w+)'),                                  F, True,  False),
    (re.compile(r'\b(\d+\.?\d*)\b'),                          N, False, False),
    (re.compile(r'\b([A-Za-z_]\w*)\b'),                       None, False, False),
]


def tokenize_line(line):
    if not line.strip():
        return [(line, BLK, False, False)]
    segs, pos, n = [], 0, len(line)
    while pos < n:
        best = None
        for pat, col, bold, italic in TOKENS:
            m = pat.match(line, pos)
            if m and (best is None or m.start() < best[0].start()):
                best = (m, col, bold, italic)
        if best is None:
            segs.append((line[pos], BLK, False, False)); pos += 1
            continue
        m, col, bold, italic = best
        if m.start() > pos:
            segs.append((line[pos:m.start()], BLK, False, False))
        text = m.group(0)
        if col is None:
            w = m.group(1)
            if w in PY_KW:
                segs.append((text, K, True, False))
            elif w in PY_BI:
                segs.append((text, F, False, False))
            else:
                segs.append((text, BLK, False, False))
        else:
            segs.append((text, col, bold, italic))
        pos = m.end()
    return segs


def write_lines_to_cell(cell, lines, size=8.5):
    """写入单元格：每行一个 paragraph；行距 1.5；def/class 前插入空段。"""
    first = True
    prev_blank = True
    for line in lines:
        is_func_start = bool(re.match(r'^(def |class |@)', line))
        if first:
            p = cell.paragraphs[0]; p.clear(); first = False
        elif is_func_start and not prev_blank:
            cell.add_paragraph()  # spacer
            p = cell.add_paragraph()
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.5
        prev_blank = (not line.strip())
        if not line:
            continue
        for txt, col, bold, italic in tokenize_line(line):
            if not txt:
                continue
            add_run(p, txt, col, size=size, bold=bold, italic=italic)


def set_table_borders(tbl, sz=12, color="666666"):
    tblPr = tbl._element.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement('w:{0}'.format(edge))
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(sz))
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        borders.append(b)
    tblPr.append(borders)


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


# ---------- 主流程 ----------
doc = Document()
for sec in doc.sections:
    sec.page_width = Cm(21); sec.page_height = Cm(29.7)
    sec.left_margin = Cm(1.8); sec.right_margin = Cm(1.8)
    sec.top_margin  = Cm(1.2); sec.bottom_margin = Cm(1.2)

style = doc.styles['Normal']
style.font.size = Pt(10.5)

first = True
for fname, desc in FILES:
    fpath = os.path.join(BASE, fname)
    if not os.path.exists(fpath):
        print('SKIP (not found):', fpath); continue
    with open(fpath, 'r', encoding='utf-8') as f:
        lines = clean(f.readlines())

    if not first:
        sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(8)
    first = False

    tp = doc.add_paragraph()
    tp.paragraph_format.space_before = Pt(6)
    tp.paragraph_format.space_after = Pt(4)
    r1 = tp.add_run(fname); r1.bold = True; r1.font.size = Pt(12); r1.font.name = 'Consolas'
    r2 = tp.add_run('  —  ' + desc); r2.font.size = Pt(9.5); r2.font.name = '宋体'
    r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.style = 'Table Grid'
    set_table_borders(tbl, sz=12, color="666666")
    cell = tbl.cell(0, 0)
    cell.width = Cm(17.4)
    shade_cell(cell, 'F8F8F8')
    write_lines_to_cell(cell, lines, size=8.5)
    print('{0}: {1} lines (after clean)'.format(fname, len(lines)))

doc.save(OUT)
print('\nSaved:', OUT)